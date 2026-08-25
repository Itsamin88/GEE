"""Synthetic content for the golden pilot communities.

These are SOFTWARE TEST CASES. The two pilot names are used so the fixture
exercises realistic shapes — a French site with an abandoned domain, a Dutch
site with a foundation and a planning permit — but every word below is
invented for testing. No value here is research evidence, and runs against this
fixture are stamped ``provenance_mode: FIXTURE`` (decision DCR-D022).
"""

from __future__ import annotations

import io
import zipfile
from typing import Any

# --------------------------------------------------------------------------
# Small binary artefacts, generated rather than committed
# --------------------------------------------------------------------------
def make_pdf(lines: list[str], title: str = "Rapport") -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    pdf.setTitle(title)
    y = 780
    for line in lines:
        pdf.drawString(60, y, line[:110])
        y -= 18
        if y < 60:
            pdf.showPage()
            y = 780
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def make_xlsx(sheets: dict[str, list[list[Any]]]) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    workbook.remove(workbook.active)
    for name, rows in sheets.items():
        sheet = workbook.create_sheet(name[:31])
        for row in rows:
            sheet.append(row)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def make_png(width: int, height: int, colour: tuple[int, int, int] = (90, 140, 80)) -> bytes:
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (width, height), colour)
    draw = ImageDraw.Draw(image)
    for offset in range(0, width, max(12, width // 12)):
        draw.line([(offset, 0), (offset, height)], fill=(60, 100, 55), width=2)
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def make_docx(paragraphs: list[str], headings: list[str] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for heading in headings or []:
        document.add_heading(heading, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_zip(members: dict[str, bytes]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, payload in members.items():
            archive.writestr(name, payload)
    return buffer.getvalue()


# --------------------------------------------------------------------------
# Page bodies
# --------------------------------------------------------------------------
POURGUES_FOOTER = """
<a href="http://facebook.test/pourgues">Facebook</a>
<a href="http://ancien-pourgues.test/">Notre ancien site</a>
<a href="http://annuaire.test/lieux/pourgues">Annuaire des ecolieux</a>
<a href="http://webdesigner.test/">Site realise par WebDesigner</a>
"""

BOEKEL_FOOTER = """
<a href="http://oud-boekel.test/">Ons oude website</a>
<a href="http://gemeente.test/">Gemeente Boekel</a>
"""


def page(title: str, body: str, *, lang: str = "fr", published: str | None = None,
         extra_head: str = "", footer: str | None = None) -> str:
    meta = f'<meta property="article:published_time" content="{published}">' if published else ""
    footer_html = POURGUES_FOOTER if footer is None else footer
    return f"""<!DOCTYPE html>
<html lang="{lang}"><head><meta charset="utf-8"><title>{title}</title>{meta}{extra_head}
<link rel="alternate" type="application/rss+xml" href="/feed">
<script src="/wp-content/themes/site.js"></script></head>
<body>
<nav><a href="/">Accueil</a><a href="/histoire">Notre histoire</a><a href="/le-lieu">Le lieu</a>
<a href="/projets">Projets</a><a href="/documents">Documents</a><a href="/blog">Blog</a></nav>
<main>{body}</main>
<footer>{footer_html}</footer></body></html>"""


POURGUES_HOME = page(
    "EcoVillage de Pourgues",
    """<h1>EcoVillage de Pourgues</h1>
<p>EcoVillage de Pourgues est un ecolieu fonde en 2015 par Ael Buffet, en Ariege.
L'association Pourgues Vivant gere le domaine.</p>
<p>Notre objectif est de regenerer les sols et la biodiversite du lieu.</p>
<p>Nous sommes 40 habitants permanents en 2023. Nous accueillons aussi 800 visiteurs par an.</p>
<p>Le domaine de 55 hectares comprend une foret, des prairies et des zones cultivees.
Nous cultivons environ 4 hectares en maraichage depuis 2018.</p>
<p>Membre du reseau GEN Europe.</p>
<img src="/img/logo-header.png" alt="logo">
<img src="/img/plan-de-masse-2016.png" alt="plan de masse">
<figure><img src="/img/foret-jardin-2019.png" alt="foret jardin">
<figcaption>La foret-jardin plantee en 2016, photographiee en 2019</figcaption></figure>""",
    published="2023-05-04",
)

POURGUES_HISTORY = page(
    "Notre histoire — EcoVillage de Pourgues",
    """<h1>Notre histoire</h1>
<p>L'association a ete creee en 2015. Le terrain a ete achete en 2015 apres deux ans de recherche.</p>
<p>Les premiers habitants se sont installes en 2016.</p>
<p>Nous avons plante 3000 arbres en 2016 pour la foret-jardin, sur une surface de 1,8 hectare.</p>
<p>En 2017 nous avons creuse des baissieres sur les pentes pour retenir l'eau de pluie.</p>
<p>Depuis 2018 nous pratiquons le paillage permanent sur toutes les planches de culture, et
nous ne labourons plus.</p>
<p>En 2019 nous avons plante encore 500 arbres dans la foret-jardin.</p>
<p>Coordonnees du lieu: 43.0561, 1.8342. Le lieu est situe a 450 m d'altitude.</p>""",
    published="2021-03-11",
)

POURGUES_LAND = page(
    "Le lieu — EcoVillage de Pourgues",
    """<h1>Le lieu</h1>
<p>Le domaine de 55 hectares est detenu collectivement par l'association.
La propriete est d'un seul tenant.</p>
<p>Nous cultivons 4 hectares en maraichage biologique certifie Ecocert.
Nous n'utilisons aucun pesticide de synthese.</p>
<p>Un systeme de recuperation d'eau de pluie alimente les cuves du jardin.</p>
<p>Nous n'irriguons pas les prairies; elles sont entierement en pluvial.</p>
<figure><img src="/img/carte-usage-des-sols.png" alt="carte usage des sols">
<figcaption>Carte d'usage des sols, 2020</figcaption></figure>""",
    published="2022-09-01",
)

POURGUES_PROJECTS = page(
    "Projets — EcoVillage de Pourgues",
    """<h1>Projets</h1>
<p>Projet de restauration de la prairie humide, commence en 2017, avec creation d'une mare.</p>
<p>Plantation de haies bocageres en bordure des parcelles depuis 2018.</p>
<p>Nous accueillons des benevoles, des stages et des formations toute l'annee.</p>
<p>Ce projet a recu un financement du programme LEADER en 2017.</p>""",
    published="2020-06-15",
)

POURGUES_DOCUMENTS = page(
    "Documents — EcoVillage de Pourgues",
    """<h1>Documents</h1>
<ul>
<li><a href="/docs/rapport-annuel-2019.pdf">Rapport annuel 2019</a></li>
<li><a href="/docs/inventaire-plantations.xlsx">Inventaire des plantations</a></li>
<li><a href="/docs/plan-de-gestion-2018.docx">Plan de gestion 2018</a></li>
<li><a href="/docs/dossier-2017.zip">Dossier de projet 2017</a></li>
<li><a href="/docs/corrompu.pdf">Rapport 2015 (fichier abime)</a></li>
<li><a href="/docs/mystere">Document sans extension</a></li>
</ul>""",
    published="2020-01-20",
)

POURGUES_BLOG = page(
    "Blog — EcoVillage de Pourgues",
    """<h1>Blog</h1>
<ul>
<li><a href="/blog/2016/plantation">Plantation de la foret-jardin (2016)</a></li>
<li><a href="/blog/2017/baissieres">Chantier baissieres (2017)</a></li>
<li><a href="/blog/2021/bilan">Bilan 2021</a></li>
</ul>
<a href="/blog/page/2" rel="next">Articles plus anciens</a>""",
    published="2021-12-01",
)

POURGUES_BLOG_2016 = page(
    "Plantation de la foret-jardin",
    """<h1>Plantation de la foret-jardin</h1>
<p>En mars 2016 nous avons plante 3000 arbres sur 1,8 hectare: pommiers, noisetiers,
chataigniers et un etage arbustif complet.</p>
<p>C'est notre premiere intervention ecologique sur le terrain.</p>""",
    published="2016-03-22",
)

POURGUES_BLOG_2017 = page(
    "Chantier baissieres",
    """<h1>Chantier baissieres</h1>
<p>En 2017 nous avons creuse 400 metres de baissieres sur courbe de niveau,
pour ralentir et infiltrer l'eau de pluie.</p>""",
    published="2017-10-02",
)

POURGUES_BLOG_2021 = page(
    "Bilan 2021",
    """<h1>Bilan 2021</h1>
<p>Nous cultivons maintenant 6 hectares en maraichage.</p>
<p>Le semis d'engrais verts est systematique entre deux cultures.</p>""",
    published="2021-11-30",
)

# The former domain: older, thinner, and the only place a 2016 figure survives.
ANCIEN_HOME = page(
    "Pourgues — ancien site",
    """<h1>Pourgues</h1>
<p>Ce site n'est plus mis a jour. Rendez-vous sur notre nouveau site.</p>
<p>Nous cultivons 2 hectares en maraichage depuis 2016.</p>
<p>Le terrain de 55 hectares a ete achete en 2015.</p>
<p><a href="/docs/bulletin-2016.pdf">Bulletin 2016</a></p>""",
    published="2016-08-01",
)

# A directory listing whose text was submitted from the website: SAME group.
ANNUAIRE_LISTING = page(
    "Pourgues — Annuaire des ecolieux",
    """<h1>EcoVillage de Pourgues</h1>
<p>EcoVillage de Pourgues est un ecolieu fonde en 2015 par Ael Buffet, en Ariege.
L'association Pourgues Vivant gere le domaine.</p>
<p>Notre objectif est de regenerer les sols et la biodiversite du lieu.</p>
<p>Nous sommes 40 habitants permanents en 2023.</p>
<p>Le domaine de 55 hectares comprend une foret, des prairies et des zones cultivees.
Nous cultivons environ 4 hectares en maraichage depuis 2018.</p>
<p>Fiche creee en 2016. Derniere mise a jour 2024.</p>""",
    published="2016-04-02",
)

# The Dutch pilot: a foundation, a permit, and a planting record.
BOEKEL_HOME = page(
    "Ecodorp Boekel",
    """<h1>Ecodorp Boekel</h1>
<p>Stichting Ecodorp Boekel is opgericht in 2014 door Ad Vlems.</p>
<p>Ons terrein beslaat 1,5 hectare in de gemeente Boekel.</p>
<p>Wij beheren 0,8 hectare in cultuur als moestuin en voedselbos.</p>
<p>In 2017 werden 1200 bomen geplant voor het voedselbos.</p>
<p>Er wonen 36 vaste bewoners.</p>
<p>Wij gebruiken geen kunstmest en geen bestrijdingsmiddelen.</p>
<figure><img src="/img/inrichtingsplan-2016.png" alt="inrichtingsplan">
<figcaption>Inrichtingsplan van het terrein, 2016</figcaption></figure>
<p><a href="/docs/omgevingsvergunning-2016.pdf">Omgevingsvergunning 2016</a></p>""",
    lang="nl",
    published="2022-04-12",
    footer=BOEKEL_FOOTER,
)

BOEKEL_HISTORY = page(
    "Geschiedenis — Ecodorp Boekel",
    """<h1>Geschiedenis</h1>
<p>De stichting is opgericht in 2014. De grond is in 2015 aangekocht van de gemeente.</p>
<p>De eerste bewoners zijn in 2018 ingetrokken.</p>
<p>In 2016 is de wadi aangelegd voor waterberging en in 2017 zijn 1200 bomen geplant.</p>
<p>Sinds 2019 passen wij niet-kerende grondbewerking toe op alle bedden.</p>""",
    lang="nl",
    published="2021-07-08",
    footer=BOEKEL_FOOTER,
)


OUD_BOEKEL_HOME = page(
    "Ecodorp Boekel — oude website",
    """<h1>Ecodorp Boekel</h1>
<p>Deze website wordt niet meer bijgewerkt.</p>
<p>In 2016 is de wadi aangelegd voor waterberging op het terrein van 1,5 hectare.</p>
<p>Wij beheren 0,5 hectare in cultuur.</p>""",
    lang="nl",
    published="2016-11-03",
    footer=BOEKEL_FOOTER,
)
